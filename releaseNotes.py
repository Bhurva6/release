import base64
import requests
import os
import streamlit as st
import pandas as pd
from dotenv import load_dotenv
from docx import Document
from io import BytesIO
import html

# Load environment variables
load_dotenv()

# Azure DevOps configuration
organization = "BFLDevOpsOrg"
project = "Martech_Unit_Projects"
repository = "3in1cms"  # Verify the exact repository name or use repository ID if necessary
pat = os.getenv("AZURE_DEVOPS_PAT")

if not pat:
    st.error("PAT token not found. Please set AZURE_DEVOPS_PAT in your .env file.")
    st.stop()

# Create proper base64 encoded authorization header
auth_string = base64.b64encode(f":{pat}".encode()).decode()

# Set headers with correct authentication
headers = {
    "Authorization": f"Basic {auth_string}",
    "Accept": "application/json"
}

# Helper function to get pull requests merged into a specified branch
def get_pull_requests(repo_id, branch_name):
    url_prs = (f"https://dev.azure.com/{organization}/{project}/_apis/git/repositories/{repo_id}/pullrequests"
               f"?searchCriteria.targetRefName=refs/heads/{branch_name}&searchCriteria.status=completed&api-version=6.0")
    response_prs = requests.get(url_prs, headers=headers)
    response_prs.raise_for_status()
    return response_prs.json()["value"]

# Helper function to get work items linked to a pull request
def get_work_items_from_pr(repo_id, pr_id):
    url_work_items = (f"https://dev.azure.com/{organization}/{project}/_apis/git/repositories/{repo_id}/pullRequests/{pr_id}/workItems?api-version=6.0")
    response_work_items = requests.get(url_work_items, headers=headers)
    
    if response_work_items.status_code == 404:
        return []
    response_work_items.raise_for_status()
    return response_work_items.json()["value"]

# Main function to generate release notes
def generate_release_notes(branch_name):
    pull_requests = get_pull_requests(repository, branch_name)
    
    user_stories = []
    bugs = []

    for pr in pull_requests:
        work_items = get_work_items_from_pr(repository, pr["pullRequestId"])
        
        for work_item in work_items:
            work_item_url = f"https://dev.azure.com/{organization}/{project}/_apis/wit/workitems/{work_item['id']}?api-version=6.0"
            work_item_response = requests.get(work_item_url, headers=headers)
            work_item_details = work_item_response.json()
            
            work_item_type = work_item_details['fields']['System.WorkItemType']
            work_item_title = work_item_details['fields']['System.Title']
            
            # Azure URLs for work items and PRs
            work_item_link = f"https://dev.azure.com/{organization}/{project}/_workitems/edit/{work_item['id']}"
            pr_link = f"https://dev.azure.com/{organization}/{project}/_git/{repository}/pullrequest/{pr['pullRequestId']}"

            # Embed links as HTML
            work_item_html = f"<a href='{work_item_link}' target='_blank'>{work_item_title}</a>"
            pr_html = f"<a href='{pr_link}' target='_blank'>{pr['title']}</a>"

            item_info = {
                'id': work_item['id'],
                'title': work_item_html,
                'pr_id': pr['pullRequestId'],
                'pr_title': pr_html
            }
            
            if work_item_type == 'User Story':
                user_stories.append(item_info)
            elif work_item_type == 'Bug':
                bugs.append(item_info)

    # Create User Stories and Bugs tables with HTML links
    user_stories_data = []
    for story in user_stories:
        user_stories_data.append([story['title'], story['pr_title']])

    bugs_data = []
    for bug in bugs:
        bugs_data.append([bug['title'], bug['pr_title']])

    # Convert to DataFrame
    user_stories_df = pd.DataFrame(user_stories_data, columns=["User Story", "PR"])
    bugs_df = pd.DataFrame(bugs_data, columns=["Bug", "PR"])

    # Adjust the index to start from 1
    user_stories_df.index = range(1, len(user_stories_df) + 1)
    bugs_df.index = range(1, len(bugs_df) + 1)

    return user_stories_df, bugs_df

# Function to generate a Word document with plain text (no HTML)
import html

# Function to generate a Word document with plain text (no HTML)
def generate_word_file(user_stories_df, bugs_df):
    doc = Document()
    doc.add_heading("Release Notes", 0)
    
    # Helper function to extract text from HTML link
    def extract_text(html_link):
        # Extract text between > and < from the HTML link
        start = html_link.find('>') + 1
        end = html_link.find('<', start)
        return html_link[start:end] if start > 0 and end > start else html_link

    # User Stories Section
    doc.add_heading('User Stories', level=1)
    if not user_stories_df.empty:
        # Create a table for user stories
        user_stories_table = doc.add_table(rows=len(user_stories_df) + 1, cols=2)
        user_stories_table.style = 'Table Grid'
        
        # Header row
        user_stories_table.cell(0, 0).text = 'User Story'
        user_stories_table.cell(0, 1).text = 'Pull Request'
        
        # Populate table rows
        for idx, row in user_stories_df.iterrows():
            user_story_text = extract_text(row['User Story'])
            pr_text = extract_text(row['PR'])
            
            user_stories_table.cell(idx, 0).text = user_story_text
            user_stories_table.cell(idx, 1).text = pr_text
    else:
        doc.add_paragraph("No user stories found.")
    
    # Bugs Section
    doc.add_heading('Bugs', level=1)
    if not bugs_df.empty:
        # Create a table for bugs
        bugs_table = doc.add_table(rows=len(bugs_df) + 1, cols=2)
        bugs_table.style = 'Table Grid'
        
        # Header row
        bugs_table.cell(0, 0).text = 'Bug'
        bugs_table.cell(0, 1).text = 'Pull Request'
        
        # Populate table rows
        for idx, row in bugs_df.iterrows():
            bug_text = extract_text(row['Bug'])
            pr_text = extract_text(row['PR'])
            
            bugs_table.cell(idx, 0).text = bug_text
            bugs_table.cell(idx, 1).text = pr_text
    else:
        doc.add_paragraph("No bugs found.")
    
    # Save document to memory as BytesIO object
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io


# Add CSS for black text area and left-align column headers
# Add CSS for black text area, left-align column headers, and copy button
st.markdown("""
    <style>
        .output-container {
            border: 1px solid #333;
            padding: 15px;
            border-radius: 5px;
            background-color: #000;  /* Black background */
            color: #fff;  /* White text */
            position: relative;
        }
        .copy-button {
            position: absolute;
            top: 10px;
            right: 10px;
            background-color: #4CAF50;
            color: white;
            border: none;
            padding: 5px 10px;
            cursor: pointer;
            border-radius: 4px;
        }
        .copy-button:hover {
            background-color: #45a049;
        }
        th {
            text-align: left !important;
        }
        td {
            color: #fff;  /* Ensure table cell text is also white */
        }
        h4 {
            margin-top: 0;  /* Remove top margin for titles inside the container */
        }
    </style>
    <script>
    function copyToClipboard() {
        // Create a temporary textarea to hold the text
        var tempTextArea = document.createElement('textarea');
        tempTextArea.value = document.querySelector('.output-container').innerText;
        document.body.appendChild(tempTextArea);
        
        // Select the text
        tempTextArea.select();
        tempTextArea.setSelectionRange(0, 99999); // For mobile devices
        
        // Copy the text
        document.execCommand('copy');
        
        // Remove the temporary textarea
        document.body.removeChild(tempTextArea);
        
        // Optional: Change button text briefly to show copied
        var copyButton = document.querySelector('.copy-button');
        var originalText = copyButton.innerText;
        copyButton.innerText = 'Copied!';
        setTimeout(() => {
            copyButton.innerText = originalText;
        }, 2000);
    }
    </script>
""", unsafe_allow_html=True)

# In the Streamlit markdown section, modify to include copy button
# Streamlit UI
st.title("Release Notes")

branch_name = st.text_input("Enter the branch name", value="")

# Place the download button above the output container
if st.button("Make Release Notes"):
    try:
        user_stories_df, bugs_df = generate_release_notes(branch_name)

        # Provide download button for the generated Word file
        if not user_stories_df.empty or not bugs_df.empty:
            word_file = generate_word_file(user_stories_df, bugs_df)
            st.download_button(
                label="Download Release Notes as Word Document",
                data=word_file,
                file_name="release_notes.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

            # Add CSS for black text area, left-align column headers, and copy button
            st.markdown("""
                <style>
                    .output-container {
                        border: 1px solid #333;
                        padding: 15px;
                        border-radius: 5px;
                        background-color: #000;  /* Black background */
                        color: #fff;  /* White text */
                        position: relative;
                    }
                    .copy-button {
                        position: absolute;
                        top: 10px;
                        right: 10px;
                        background-color: #4CAF50;
                        color: white;
                        border: none;
                        padding: 5px 10px;
                        cursor: pointer;
                        border-radius: 4px;
                    }
                    .copy-button:hover {
                        background-color: #45a049;
                    }
                    th {
                        text-align: left !important;
                    }
                    td {
                        color: #fff;  /* Ensure table cell text is also white */
                    }
                    h4 {
                        margin-top: 0;  /* Remove top margin for titles inside the container */
                    }
                </style>
                <script>
                function copyToClipboard() {
                    var tempTextArea = document.createElement('textarea');
                    tempTextArea.value = document.querySelector('.output-container').innerText;
                    document.body.appendChild(tempTextArea);
                    
                    tempTextArea.select();
                    tempTextArea.setSelectionRange(0, 99999);
                    
                    document.execCommand('copy');
                    
                    document.body.removeChild(tempTextArea);
                    
                    var copyButton = document.querySelector('.copy-button');
                    var originalText = copyButton.innerText;
                    copyButton.innerText = 'Copied!';
                    setTimeout(() => {
                        copyButton.innerText = originalText;
                    }, 2000);
                }
                </script>
            """, unsafe_allow_html=True)

            # Display User Stories and Bugs inside a black text area
            st.markdown(
                f"""
                <div class="output-container">
                    <button onclick="copyToClipboard()" class="copy-button">Copy</button>
                    <h4>User Stories</h4>
                    {user_stories_df.to_html(escape=False, index=True)}
                    <h4>Bugs</h4>
                    {bugs_df.to_html(escape=False, index=True)}
                </div>
                """,
                unsafe_allow_html=True
            )
        else:
            st.warning("No release notes found for the specified branch.")

    except requests.exceptions.HTTPError as e:
        st.error(f"Failed to fetch data: {e}")